#!/usr/bin/env python3
"""
DOCX fájlok validálása
"""
import os
from pathlib import Path

def validate_docx_files():
    """DOCX fájlok validálása"""
    docx_dir = "documentation_docx"
    
    if not os.path.exists(docx_dir):
        print("❌ documentation_docx könyvtár nem található!")
        return
    
    docx_files = list(Path(docx_dir).glob("*.docx"))
    
    print("📚 DOCX FÁJLOK VALIDÁLÁSA")
    print("=" * 30)
    
    for docx_file in sorted(docx_files):
        try:
            file_size = docx_file.stat().st_size
            size_kb = round(file_size / 1024, 2)
            
            # Zip fájl ellenőrzés (DOCX valójában zip)
            with open(docx_file, 'rb') as f:
                header = f.read(4)
                is_zip = header == b'PK\x03\x04'
            
            status = "✅ Érvényes" if is_zip else "❌ Sérült"
            print(f"{status} {docx_file.name} ({size_kb} KB)")
            
            # Python-docx próba
            try:
                from docx import Document
                doc = Document(str(docx_file))
                para_count = len(doc.paragraphs)
                print(f"   📄 Bekezdések: {para_count}")
                if para_count > 0:
                    first_text = doc.paragraphs[0].text[:40]
                    print(f"   📝 Első sor: {first_text}...")
            except Exception as e:
                print(f"   ⚠️ Python-docx hiba: {e}")
                
        except Exception as e:
            print(f"❌ Hiba {docx_file.name}: {e}")
    
    print("\n🎯 Validálás befejezve!")

if __name__ == "__main__":
    validate_docx_files()
