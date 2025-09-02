#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX konverter
Konvertálja a projekt markdown dokumentációit DOCX formátumba
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_hyperlink(paragraph, text, url):
    """Hyperlink létrehozása a dokumentumban"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Hyperlink stílus (kék szín, aláhúzás)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def setup_document_styles(doc):
    """Dokumentum stílusok beállítása"""
    styles = doc.styles
    
    # Címsor stílusok
    try:
        heading1 = styles['Heading 1']
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1.font.color.rgb = None  # Alapértelmezett szín
    except:
        pass
    
    try:
        heading2 = styles['Heading 2']
        heading2.font.size = Pt(16)
        heading2.font.bold = True
    except:
        pass
    
    try:
        heading3 = styles['Heading 3']
        heading3.font.size = Pt(14)
        heading3.font.bold = True
    except:
        pass

def convert_markdown_to_docx(md_file_path, output_dir):
    """Markdown fájl konvertálása DOCX formátumba"""
    print(f"🔄 Konvertálás: {md_file_path}")  # [DEBUG] commented out for production
    
    # Markdown fájl beolvasása
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # DOCX dokumentum létrehozása
    doc = Document()
    setup_document_styles(doc)
    
    # Fájl nevből cím generálása
    file_name = Path(md_file_path).stem
    title_map = {
        'README': '📋 Plágium Ellenőrző Rendszer - Áttekintés',
        'DOKUMENTACIO': '📚 Plágium Ellenőrző Rendszer - Teljes Dokumentáció', 
        'TECHNIKAI_ATTEKINTES': '🔧 Plágium Ellenőrző Rendszer - Technikai Áttekintés',
        'GYORS_REFERENCIA': '⚡ Plágium Ellenőrző Rendszer - Gyors Referencia'
    }
    
    main_title = title_map.get(file_name, f"📄 {file_name}")
    
    # Főcím hozzáadása
    title_para = doc.add_heading(main_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Üres sor
    doc.add_paragraph()
    
    # Markdown tartalom feldolgozása soronként
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Kódblokk kezelése
        if line.startswith('```'):
            if in_code_block:
                # Kódblokk vége
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    code_para = doc.add_paragraph(code_text)
                    code_para.style = 'No Spacing'
                    # Monospace font
                    for run in code_para.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                in_code_block = False
                code_lines = []
            else:
                # Kódblokk kezdete
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Táblázat kezelése
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and line.strip() == '':
            # Táblázat vége
            if table_lines:
                create_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
            doc.add_paragraph()  # Üres sor a táblázat után
            i += 1
            continue
        elif in_table:
            # Táblázat vége (nem üres sor)
            if table_lines:
                create_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
            # Folytatjuk az aktuális sorral
        
        # Címsorok
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            if level <= 3:  # Csak 3 szintig
                title_text = line.lstrip('#').strip()
                # Emoji eltávolítása a címből (opcionális)
                title_text = re.sub(r'^[^\w\s]+\s*', '', title_text)
                doc.add_heading(title_text, level=level)
            i += 1
            continue
        
        # Üres sorok
        if line.strip() == '':
            doc.add_paragraph()
            i += 1
            continue
        
        # Lista elemek
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
            list_text = line[2:].strip() if line.startswith(('- ', '* ')) else re.sub(r'^\d+\.\s*', '', line)
            list_text = clean_markdown_formatting(list_text)
            para = doc.add_paragraph(list_text, style='List Bullet' if line.startswith(('- ', '* ')) else 'List Number')
            i += 1
            continue
        
        # Normál bekezdések
        if line.strip():
            clean_text = clean_markdown_formatting(line)
            para = doc.add_paragraph(clean_text)
            i += 1
            continue
        
        i += 1
    
    # Ha a végén maradt táblázat
    if in_table and table_lines:
        create_table_from_markdown(doc, table_lines)
    
    # Kimeneti fájl neve
    output_file = os.path.join(output_dir, f"{file_name}.docx")
    
    # Dokumentum mentése
    doc.save(output_file)
    print(f"✅ Elkészült: {output_file}")  # [DEBUG] commented out for production
    return output_file

def create_table_from_markdown(doc, table_lines):
    """Markdown táblázat konvertálása DOCX táblázattá"""
    if len(table_lines) < 2:
        return
    
    # Fejléc sor feldolgozása
    header_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Elválasztó sor kihagyása (table_lines[1])
    
    # Adat sorok feldolgozása
    data_rows = []
    for line in table_lines[2:]:
        if line.strip():
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            data_rows.append(row)
    
    if not data_rows:
        return
    
    # Táblázat létrehozása
    table = doc.add_table(rows=1, cols=len(header_row))
    table.style = 'Table Grid'
    
    # Fejléc kitöltése
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(header_row):
        if i < len(header_cells):
            header_cells[i].text = clean_markdown_formatting(header_text)
            # Fejléc formázás
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Adat sorok hozzáadása
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = clean_markdown_formatting(cell_text)

def clean_markdown_formatting(text):
    """Markdown formázás eltávolítása/konvertálása"""
    # Félkövér (**text** vagy __text__)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    
    # Dőlt (*text* vagy _text_)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    
    # Kód (`code`)
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Linkek [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # HTML tagek eltávolítása
    text = re.sub(r'<[^>]+>', '', text)
    
    return text

def main():
    """Főprogram - az összes MD fájl konvertálása"""
    print("📚 MARKDOWN TO DOCX KONVERTER")  # [DEBUG] commented out for production
    print("=" * 40)  # [DEBUG] commented out for production
    
    # Aktuális könyvtár
    current_dir = os.getcwd()
    
    # Kimeneti könyvtár
    output_dir = os.path.join(current_dir, "documentation_docx")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    print(f"📁 Kimeneti könyvtár létrehozva: {output_dir}")  # [DEBUG] commented out for production
    
    # Markdown fájlok keresése
    md_files = [f for f in os.listdir(current_dir) if f.endswith('.md')]
    
    if not md_files:
        print("❌ Nem találhatók markdown fájlok az aktuális könyvtárban!")  # [DEBUG] commented out for production
        return
    
    print(f"📋 Talált MD fájlok: {len(md_files)}")  # [DEBUG] commented out for production
    print()  # [DEBUG] commented out for production
    
    # Konvertálás
    converted_files = []
    for md_file in sorted(md_files):
        try:
            output_file = convert_markdown_to_docx(md_file, output_dir)
            converted_files.append(output_file)
        except Exception as e:
            print(f"❌ Hiba a {md_file} konvertálásakor: {e}")  # [DEBUG] commented out for production
    
    print()  # [DEBUG] commented out for production
    print("=" * 40)  # [DEBUG] commented out for production
    print(f"✅ KONVERTÁLÁS BEFEJEZVE!")  # [DEBUG] commented out for production
    print(f"📊 Sikeresen konvertált fájlok: {len(converted_files)}")  # [DEBUG] commented out for production
    print(f"📁 Kimeneti könyvtár: {output_dir}")  # [DEBUG] commented out for production
    print()  # [DEBUG] commented out for production
    
    for file_path in converted_files:
        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        size_kb = round(file_size / 1024, 2)
    print(f"  📄 {file_name} ({size_kb} KB)")  # [DEBUG] commented out for production
    
    print()  # [DEBUG] commented out for production
    print("🎯 A DOCX fájlok készen állnak!")  # [DEBUG] commented out for production

if __name__ == "__main__":
    main()
