from docx import Document
import re

def read_docx(file_path):
    try:
        #print(f"Beolvasom a fájlt: {file_path}")
        doc = Document(file_path)
        full_text = []
        #táblák olvasása
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text+". \n")
        result = "\n".join(full_text)
        
        #paragrafusok olvasása
        for para in doc.paragraphs:
            if para.text.strip():  # Csak ha nem üres a sor
                full_text.append(para.text)
        result = "\n".join(full_text)
        
        # Normalizálás: en dash, em dash -> sima kötőjel
        #result = re.sub(r'[\u2010-\u2015\u2212]', '-', result)
        # :- csere ! re
        # result = result.replace('-', '!')
        # CR LF csere <br> re
        result = result.replace('\r\n', '').replace('\n', '').replace('\r', '')
        # Csere: ? szóköz és nagybetű után <br>
        result = re.sub(r'\?\s*(?=[A-ZÁÉÍÓÖŐÚÜŰ])', '?\n', result)
        result = re.sub(r'\!\s*(?=[A-ZÁÉÍÓÖŐÚÜŰ])', '!\n', result)
        result = re.sub(r'\.\s*(?=[A-ZÁÉÍÓÖŐÚÜŰ])', '.\n', result)
        return result
    except Exception as e:
        print(f"[HIBA] DOCX beolvasás sikertelen: {file_path} ({e})")
        return f"[HIBA] DOCX beolvasás sikertelen: {file_path} ({e})"

